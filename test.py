import docx

document = docx.Document()

document.add_heading('\n\n\n\n\n\n\n\n\n\nTitle', 1).alignment = docx.\
        enum.text.WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph('Things and such \n')
p.add_run('another thing').bold = True
document.add_page_break()

document.save('thing.docx')
